#!/usr/bin/env python3
"""
Generate a comprehensive .docx file from the Java Spring Complete Course seed data.

Reads:
  - assets/seed-data/modules.json
  - assets/seed-data/lessons.json
  - assets/seed-data/keywords.json
  - assets/seed-data/quizzes.json
  - assets/seed-data/quiz_questions.json
  - assets/seed-data/code_examples.json
  - assets/seed-data/keyword_relations.json

Outputs:
  - doc/Java-Spring-Complete-Course.docx
"""

import json
import os
import re
import sys

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
SEED_DIR = os.path.join(PROJECT_ROOT, "assets", "seed-data")
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "doc")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Java-Spring-Complete-Course.docx")

# ---------------------------------------------------------------------------
# Helpers — load JSON
# ---------------------------------------------------------------------------

def load_json(filename: str) -> list:
    path = os.path.join(SEED_DIR, filename)
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# Helpers — formatting
# ---------------------------------------------------------------------------

DIFFICULTY_LABELS = {
    "beginner": "🟢 Beginner",
    "intermediate": "🟡 Intermediate",
    "advanced": "🔴 Advanced",
    "expert": "🟣 Expert",
}


def set_cell_shading(cell, color_hex: str):
    """Set background shading on a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, color_hex: str):
    """Set background shading on a paragraph (for code blocks)."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    paragraph._p.get_or_add_pPr().append(shading)


def add_bold_text_with_markdown(paragraph, text: str, font_name="Calibri", font_size=Pt(11)):
    """Parse **bold** markdown in text and add runs accordingly."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.name = font_name
        run.font.size = font_size


def add_inline_code_text(paragraph, text: str, font_name="Calibri", font_size=Pt(11)):
    """Parse `code` and **bold** markdown in text and add runs accordingly."""
    # First split by inline code
    code_parts = re.split(r"(`[^`]+`)", text)
    for code_part in code_parts:
        if code_part.startswith("`") and code_part.endswith("`"):
            run = paragraph.add_run(code_part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        else:
            # Handle bold within non-code parts
            bold_parts = re.split(r"(\*\*.*?\*\*)", code_part)
            for bp in bold_parts:
                if bp.startswith("**") and bp.endswith("**"):
                    run = paragraph.add_run(bp[2:-2])
                    run.bold = True
                    run.font.name = font_name
                    run.font.size = font_size
                else:
                    run = paragraph.add_run(bp)
                    run.font.name = font_name
                    run.font.size = font_size


# ---------------------------------------------------------------------------
# Style setup
# ---------------------------------------------------------------------------

def setup_styles(doc: Document):
    """Configure document styles: margins, fonts, heading styles."""
    # Page margins — 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2C, 0x5F, 0x2D)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x6B, 0x46, 0x26)

    return doc


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------

def add_title_page(doc: Document):
    """Add a title page with course name."""
    # Add some blank lines for vertical centering
    for _ in range(6):
        doc.add_paragraph("")

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Java Spring Complete Course")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    # Vietnamese subtitle
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Khóa Học Java Spring Toàn Diện")
    run2.font.name = "Calibri"
    run2.font.size = Pt(20)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Separator
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("━" * 40)
    run3.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run3.font.size = Pt(14)

    # Description
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Từ Java Core đến Spring Cloud & Microservices")
    run4.font.name = "Calibri"
    run4.font.size = Pt(14)
    run4.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("9 Modules • 27 Lessons • Production-Ready Knowledge")
    run5.font.name = "Calibri"
    run5.font.size = Pt(12)
    run5.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Page break
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Table of Contents (manual)
# ---------------------------------------------------------------------------

def add_table_of_contents(doc: Document, modules: list, lessons_by_module: dict):
    """Add a manual table of contents listing all modules and lessons."""
    doc.add_heading("Mục Lục — Table of Contents", level=1)
    doc.add_paragraph("")

    for module in modules:
        mid = module["id"]
        order = module["order_index"]
        title = module["title"]
        title_vi = module.get("title_vi", "")

        p = doc.add_paragraph()
        run = p.add_run(f"Module {order}: {title} — {title_vi}")
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

        module_lessons = lessons_by_module.get(mid, [])
        for lesson in module_lessons:
            lorder = lesson["order_index"]
            ltitle = lesson["title"]
            ltitle_vi = lesson.get("title_vi", "")
            lp = doc.add_paragraph()
            lp.paragraph_format.left_indent = Inches(0.5)
            run2 = lp.add_run(f"Bài {lorder}: {ltitle} — {ltitle_vi}")
            run2.font.name = "Calibri"
            run2.font.size = Pt(11)
            run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Content rendering — parse content_json sections
# ---------------------------------------------------------------------------

def render_heading(doc: Document, section: dict):
    """Render a heading section."""
    level = section.get("level", 2)
    text = section.get("text", "")
    # Clamp level to 1-3 for Word headings; level 1 in content maps to H2 in doc
    # (H1 is reserved for module titles)
    doc_level = min(max(level, 1), 3)
    if doc_level == 1:
        doc_level = 2  # Promote content H1 to H2 (H1 = module)
    doc.add_heading(text, level=doc_level)


def render_paragraph(doc: Document, section: dict):
    """Render a paragraph section with bold/code markdown support."""
    text = section.get("text", "")
    if not text:
        return
    p = doc.add_paragraph()
    add_inline_code_text(p, text)


def render_code_block(doc: Document, section: dict):
    """Render a code block with monospace font and gray background."""
    code = section.get("code", "")
    if not code:
        return
    language = section.get("language", "")

    # Language label
    if language:
        label_p = doc.add_paragraph()
        run = label_p.add_run(f"  {language.upper()}")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        set_paragraph_shading(label_p, "4A4A4A")
        label_p.paragraph_format.space_after = Pt(0)
        label_p.paragraph_format.space_before = Pt(6)

    # Code content — split into lines
    lines = code.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(13)
        set_paragraph_shading(p, "F5F5F5")
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)

    # Small spacer after code block
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(4)


def render_table(doc: Document, section: dict):
    """Render a table section."""
    headers = section.get("headers", [])
    rows = section.get("rows", [])
    if not headers and not rows:
        return

    col_count = len(headers) if headers else (len(rows[0]) if rows else 0)
    if col_count == 0:
        return

    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    if headers:
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(header))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, "1A478A")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_code_text(p, str(cell_text), font_size=Pt(10))
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F4F8")

    doc.add_paragraph("")  # spacer


def render_list(doc: Document, section: dict):
    """Render a list section (ordered or unordered)."""
    items = section.get("items", [])
    ordered = section.get("ordered", False)

    for i, item in enumerate(items):
        p = doc.add_paragraph()
        if ordered:
            prefix = f"{i + 1}. "
        else:
            prefix = "• "
        add_inline_code_text(p, prefix + str(item))
        p.paragraph_format.left_indent = Inches(0.3)


def render_diagram(doc: Document, section: dict):
    """Render a diagram as a monospace code block."""
    code = section.get("code", "") or section.get("text", "")
    if not code:
        return

    # Label
    label_p = doc.add_paragraph()
    run = label_p.add_run("  📊 DIAGRAM")
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.bold = True
    set_paragraph_shading(label_p, "2C5F2D")
    label_p.paragraph_format.space_after = Pt(0)
    label_p.paragraph_format.space_before = Pt(6)

    lines = code.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(13)
        set_paragraph_shading(p, "F0FFF0")
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(4)


def render_content_sections(doc: Document, content_json_str: str):
    """Parse content_json string and render all sections."""
    if not content_json_str:
        return

    try:
        content = json.loads(content_json_str)
    except (json.JSONDecodeError, TypeError):
        return

    sections = content.get("sections", [])
    for section in sections:
        stype = section.get("type", "")
        if stype == "heading":
            render_heading(doc, section)
        elif stype == "paragraph":
            render_paragraph(doc, section)
        elif stype == "code_block":
            render_code_block(doc, section)
        elif stype == "table":
            render_table(doc, section)
        elif stype == "list":
            render_list(doc, section)
        elif stype == "diagram":
            render_diagram(doc, section)


# ---------------------------------------------------------------------------
# Quiz section
# ---------------------------------------------------------------------------

def add_quiz_section(doc: Document, lesson_id: str, quizzes: list,
                     quiz_questions_by_quiz: dict):
    """Add quiz questions for a lesson."""
    lesson_quizzes = [q for q in quizzes if q["lesson_id"] == lesson_id]
    if not lesson_quizzes:
        return

    doc.add_heading("📝 Câu Hỏi Ôn Tập", level=3)

    for quiz in lesson_quizzes:
        quiz_id = quiz["id"]
        questions = quiz_questions_by_quiz.get(quiz_id, [])
        questions.sort(key=lambda x: x.get("order_index", 0))

        for q_idx, question in enumerate(questions):
            q_text = question.get("question_text", "")
            options_raw = question.get("options_json", "[]")
            correct = question.get("correct_answer", "")
            explanation = question.get("explanation", "")

            try:
                options = json.loads(options_raw) if isinstance(options_raw, str) else options_raw
            except (json.JSONDecodeError, TypeError):
                options = []

            # Question text
            p = doc.add_paragraph()
            run = p.add_run(f"Câu {q_idx + 1}: {q_text}")
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.bold = True

            # Options
            option_labels = ["A", "B", "C", "D", "E", "F", "G", "H"]
            for o_idx, option in enumerate(options):
                label = option_labels[o_idx] if o_idx < len(option_labels) else str(o_idx + 1)
                op = doc.add_paragraph()
                op.paragraph_format.left_indent = Inches(0.3)
                opt_text = str(option)
                # Truncate very long options
                if len(opt_text) > 200:
                    opt_text = opt_text[:200] + "…"
                is_correct = (str(option) == correct) or (opt_text.startswith(correct[:50]) if correct and len(correct) > 50 else False)
                prefix = f"{'✅' if is_correct else '⬚'} {label}. "
                run = op.add_run(prefix + opt_text)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                if is_correct:
                    run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x2D)
                    run.font.bold = True

            # Explanation
            if explanation:
                ep = doc.add_paragraph()
                ep.paragraph_format.left_indent = Inches(0.3)
                set_paragraph_shading(ep, "FFF8E1")
                expl_text = explanation if len(explanation) <= 300 else explanation[:300] + "…"
                run = ep.add_run(f"💡 Giải thích: {expl_text}")
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x6B, 0x46, 0x26)

            doc.add_paragraph("")  # spacer


# ---------------------------------------------------------------------------
# Keyword summary section
# ---------------------------------------------------------------------------

def add_keyword_section(doc: Document, lesson_id: str, keywords: list):
    """Add keyword summary table for a lesson."""
    lesson_keywords = [k for k in keywords if k["lesson_id"] == lesson_id]
    if not lesson_keywords:
        return

    doc.add_heading("🔑 Tổng Hợp Keywords", level=3)

    table = doc.add_table(rows=1 + len(lesson_keywords), cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    headers = ["Keyword", "Definition", "Category"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "6B4626")

    # Data rows
    for r_idx, kw in enumerate(lesson_keywords):
        name = kw.get("name", "")
        definition = kw.get("definition", "")
        explanation = kw.get("explanation", "")
        category = kw.get("category", "")

        # Name cell
        cell0 = table.rows[r_idx + 1].cells[0]
        cell0.text = ""
        p0 = cell0.paragraphs[0]
        run0 = p0.add_run(name)
        run0.font.name = "Calibri"
        run0.font.size = Pt(10)
        run0.font.bold = True

        # Definition cell (with explanation if available)
        cell1 = table.rows[r_idx + 1].cells[1]
        cell1.text = ""
        p1 = cell1.paragraphs[0]
        def_text = definition
        if explanation:
            def_text += f"\n{explanation}" if len(explanation) <= 200 else f"\n{explanation[:200]}…"
        # Truncate very long definitions
        if len(def_text) > 400:
            def_text = def_text[:400] + "…"
        run1 = p1.add_run(def_text)
        run1.font.name = "Calibri"
        run1.font.size = Pt(9)

        # Category cell
        cell2 = table.rows[r_idx + 1].cells[2]
        cell2.text = ""
        p2 = cell2.paragraphs[0]
        run2 = p2.add_run(category)
        run2.font.name = "Calibri"
        run2.font.size = Pt(9)
        run2.font.italic = True

        if r_idx % 2 == 1:
            for c in range(3):
                set_cell_shading(table.rows[r_idx + 1].cells[c], "FFF8F0")

    doc.add_paragraph("")  # spacer


# ---------------------------------------------------------------------------
# Tech Lead Insights section
# ---------------------------------------------------------------------------

TECH_LEAD_INSIGHTS = {
    "lesson-01-01": {
        "title": "Variables & Data Types",
        "takeaways": [
            "Luôn dùng kiểu dữ liệu phù hợp nhất — int cho đếm, long cho timestamps, BigDecimal cho tiền tệ.",
            "Hiểu Stack vs Heap giúp tối ưu memory và tránh memory leaks trong production.",
            "final fields đảm bảo thread-safety qua JMM — ưu tiên immutability trong concurrent code.",
        ],
        "mistakes": [
            "Dùng float/double cho tiền tệ — luôn dùng BigDecimal.",
            "So sánh wrapper objects bằng == thay vì .equals().",
            "Quên Integer cache chỉ áp dụng cho -128 đến 127.",
        ],
        "architecture": "Kiểu dữ liệu ảnh hưởng trực tiếp đến database schema design, API contract, và serialization format. Chọn sai kiểu ở domain layer sẽ cascade ra toàn bộ hệ thống.",
        "interview": [
            "Giải thích pass-by-value trong Java — tại sao swap() không hoạt động.",
            "String Pool hoạt động thế nào và tại sao String là immutable.",
            "Autoboxing pitfalls và Integer cache.",
        ],
    },
    "lesson-01-02": {
        "title": "Operators",
        "takeaways": [
            "Short-circuit evaluation (&&, ||) không chỉ là optimization mà còn là safety mechanism.",
            "Luôn dùng Math.addExact() khi cần detect overflow trong business logic.",
            "Hiểu operator precedence tránh bugs tinh vi — khi nghi ngờ, dùng parentheses.",
        ],
        "mistakes": [
            "Nhầm lẫn & (bitwise) với && (logical) — kết quả khác nhau khi có side effects.",
            "Integer division truncation: 7/2 = 3, không phải 3.5.",
            "i = i++ không tăng i — JVM bytecode behavior.",
        ],
        "architecture": "Operator overloading không có trong Java (trừ String +). Khi design DSL hoặc fluent API, dùng method chaining thay vì cố gắng hack operators.",
        "interview": [
            "Sự khác biệt giữa == và .equals() cho objects.",
            "Short-circuit evaluation và ứng dụng trong null-safe code.",
            "Bitwise operators trong permission systems (flags).",
        ],
    },
    "lesson-01-03": {
        "title": "Control Flow",
        "takeaways": [
            "Switch expressions (Java 14+) giảm boilerplate và tránh fall-through bugs.",
            "Pattern matching (Java 16+) thay thế instanceof + cast chains.",
            "Labeled break/continue hữu ích cho nested loops nhưng dùng tiết chế.",
        ],
        "mistakes": [
            "Quên break trong switch — fall-through là nguồn bugs phổ biến.",
            "Infinite loops do điều kiện sai — luôn verify loop termination.",
            "Nested if quá sâu — refactor thành early return hoặc strategy pattern.",
        ],
        "architecture": "Control flow phức tạp là code smell. Trong production, ưu tiên polymorphism, strategy pattern, hoặc state machine thay vì switch/if chains dài.",
        "interview": [
            "Switch expression vs switch statement — khi nào dùng cái nào.",
            "Enhanced for-loop vs iterator — ConcurrentModificationException.",
            "Sealed classes và exhaustive switch (Java 17+).",
        ],
    },
    "lesson-01-04": {
        "title": "Arrays",
        "takeaways": [
            "Arrays có fixed size — dùng ArrayList khi cần dynamic sizing.",
            "System.arraycopy() nhanh hơn manual loop — JVM intrinsic.",
            "Arrays.sort() dùng dual-pivot quicksort cho primitives, TimSort cho objects.",
        ],
        "mistakes": [
            "ArrayIndexOutOfBoundsException — luôn validate index.",
            "Nhầm array reference copy với deep copy.",
            "Dùng array khi Collection phù hợp hơn.",
        ],
        "architecture": "Trong microservices, arrays thường được serialize thành JSON arrays. Cân nhắc pagination khi truyền large arrays qua network.",
        "interview": [
            "Array vs ArrayList — trade-offs về performance và flexibility.",
            "Covariance trong arrays vs generics.",
            "Multi-dimensional arrays: true 2D vs jagged arrays.",
        ],
    },
    "lesson-01-05": {
        "title": "Strings",
        "takeaways": [
            "String là immutable — mỗi concatenation tạo object mới.",
            "StringBuilder cho single-thread, StringBuffer cho multi-thread (nhưng hiếm khi cần).",
            "Text blocks (Java 13+) cho multi-line strings — clean hơn concatenation.",
        ],
        "mistakes": [
            "String concatenation trong loop — O(n²) performance.",
            "So sánh String bằng == thay vì .equals().",
            "Quên String Pool behavior khi debug.",
        ],
        "architecture": "String processing là bottleneck phổ biến. Trong high-throughput systems, cân nhắc char[] hoặc ByteBuffer cho performance-critical paths.",
        "interview": [
            "Tại sao String là immutable và final.",
            "String Pool, intern(), và memory implications.",
            "StringBuilder vs StringBuffer vs String concatenation.",
        ],
    },
    "lesson-01-06": {
        "title": "Methods",
        "takeaways": [
            "Method signature = name + parameter types. Return type KHÔNG phải part of signature.",
            "Varargs (Type...) phải là parameter cuối cùng.",
            "Static methods không cần instance — utility methods, factory methods.",
        ],
        "mistakes": [
            "Overloading confusion khi có autoboxing và varargs.",
            "Recursive methods không có base case — StackOverflowError.",
            "Side effects trong methods — khó test và debug.",
        ],
        "architecture": "Clean methods: single responsibility, meaningful names, ≤3 parameters. Dùng Builder pattern khi cần nhiều parameters.",
        "interview": [
            "Method overloading vs overriding.",
            "Pass-by-value cho primitives và references.",
            "Static vs instance methods — khi nào dùng cái nào.",
        ],
    },
    "lesson-01-07": {
        "title": "Exception Handling",
        "takeaways": [
            "Checked exceptions cho recoverable errors, unchecked cho programming errors.",
            "Try-with-resources (Java 7+) đảm bảo resource cleanup.",
            "Custom exceptions nên extend RuntimeException trong modern Java.",
        ],
        "mistakes": [
            "Catch Exception/Throwable quá rộng — mất thông tin lỗi cụ thể.",
            "Swallow exceptions (catch rỗng) — log ít nhất.",
            "Throw trong finally block — override exception gốc.",
        ],
        "architecture": "Exception hierarchy nên map với domain errors. Trong REST APIs, translate domain exceptions thành HTTP status codes qua @ControllerAdvice.",
        "interview": [
            "Checked vs unchecked exceptions — design philosophy.",
            "Try-with-resources và AutoCloseable.",
            "Exception chaining và root cause analysis.",
        ],
    },
}

# Default insights for lessons not in the map
DEFAULT_INSIGHTS = {
    "takeaways": [
        "Nắm vững concepts cơ bản trước khi áp dụng vào production code.",
        "Viết code clean, readable, và maintainable — đây là skill quan trọng nhất.",
        "Luôn viết tests cho business logic — unit tests là safety net.",
    ],
    "mistakes": [
        "Copy-paste code mà không hiểu — technical debt tích lũy.",
        "Bỏ qua error handling — production sẽ expose mọi edge case.",
        "Over-engineering — YAGNI (You Aren't Gonna Need It).",
    ],
    "architecture": "Mỗi concept đều có vị trí trong kiến trúc tổng thể. Hiểu rõ concept giúp đưa ra design decisions đúng đắn khi scale hệ thống.",
    "interview": [
        "Giải thích concept bằng ngôn ngữ đơn giản — chứng tỏ hiểu sâu.",
        "Đưa ra ví dụ thực tế từ kinh nghiệm làm việc.",
        "So sánh trade-offs giữa các approaches khác nhau.",
    ],
}


def add_tech_lead_insights(doc: Document, lesson_id: str, lesson_title: str):
    """Add Tech Lead Insights section for a lesson."""
    doc.add_heading("🧠 Mở Rộng Tư Duy Tech Lead", level=3)

    insights = TECH_LEAD_INSIGHTS.get(lesson_id, None)
    if insights is None:
        insights = dict(DEFAULT_INSIGHTS)
        insights["title"] = lesson_title

    # Key takeaways
    p_header = doc.add_paragraph()
    run = p_header.add_run("🎯 Key Takeaways cho Production Code:")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    takeaways = insights.get("takeaways", DEFAULT_INSIGHTS["takeaways"])
    for item in takeaways:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(f"✦ {item}")
        run.font.name = "Calibri"
        run.font.size = Pt(10)

    # Common mistakes
    p_header2 = doc.add_paragraph()
    run2 = p_header2.add_run("⚠️ Common Mistakes to Avoid:")
    run2.font.name = "Calibri"
    run2.font.size = Pt(11)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)

    mistakes = insights.get("mistakes", DEFAULT_INSIGHTS["mistakes"])
    for item in mistakes:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(f"✗ {item}")
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x33, 0x00)

    # Architecture connection
    p_header3 = doc.add_paragraph()
    run3 = p_header3.add_run("🏗️ Kết Nối Với Kiến Trúc Thực Tế:")
    run3.font.name = "Calibri"
    run3.font.size = Pt(11)
    run3.font.bold = True
    run3.font.color.rgb = RGBColor(0x2C, 0x5F, 0x2D)

    arch_text = insights.get("architecture", DEFAULT_INSIGHTS["architecture"])
    p_arch = doc.add_paragraph()
    p_arch.paragraph_format.left_indent = Inches(0.3)
    set_paragraph_shading(p_arch, "F0FFF0")
    run = p_arch.add_run(arch_text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)

    # Interview points
    p_header4 = doc.add_paragraph()
    run4 = p_header4.add_run("🎤 Interview-Level Understanding:")
    run4.font.name = "Calibri"
    run4.font.size = Pt(11)
    run4.font.bold = True
    run4.font.color.rgb = RGBColor(0x6B, 0x46, 0x26)

    interview = insights.get("interview", DEFAULT_INSIGHTS["interview"])
    for item in interview:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(f"💬 {item}")
        run.font.name = "Calibri"
        run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Main generation
# ---------------------------------------------------------------------------

def main():
    print("=" * 60)
    print("📚 Java Spring Complete Course — DOCX Generator")
    print("=" * 60)

    # Load all data
    print("\n📂 Loading seed data...")
    modules = load_json("modules.json")
    lessons = load_json("lessons.json")
    keywords = load_json("keywords.json")
    quizzes = load_json("quizzes.json")
    quiz_questions = load_json("quiz_questions.json")
    code_examples = load_json("code_examples.json")
    keyword_relations = load_json("keyword_relations.json")

    print(f"   ✓ {len(modules)} modules")
    print(f"   ✓ {len(lessons)} lessons")
    print(f"   ✓ {len(keywords)} keywords")
    print(f"   ✓ {len(quizzes)} quizzes")
    print(f"   ✓ {len(quiz_questions)} quiz questions")
    print(f"   ✓ {len(code_examples)} code examples")
    print(f"   ✓ {len(keyword_relations)} keyword relations")

    # Sort modules by order_index
    modules.sort(key=lambda m: m.get("order_index", 0))

    # Group lessons by module_id, sorted by order_index
    lessons_by_module = {}
    for lesson in lessons:
        mid = lesson["module_id"]
        if mid not in lessons_by_module:
            lessons_by_module[mid] = []
        lessons_by_module[mid].append(lesson)
    for mid in lessons_by_module:
        lessons_by_module[mid].sort(key=lambda l: l.get("order_index", 0))

    # Group quiz questions by quiz_id
    quiz_questions_by_quiz = {}
    for qq in quiz_questions:
        qid = qq["quiz_id"]
        if qid not in quiz_questions_by_quiz:
            quiz_questions_by_quiz[qid] = []
        quiz_questions_by_quiz[qid].append(qq)

    # Create document
    print("\n📝 Creating document...")
    doc = Document()
    setup_styles(doc)

    # Title page
    print("   → Title page")
    add_title_page(doc)

    # Table of Contents
    print("   → Table of Contents")
    add_table_of_contents(doc, modules, lessons_by_module)

    # Process each module
    total_lessons = 0
    for module in modules:
        mid = module["id"]
        order = module["order_index"]
        title = module["title"]
        title_vi = module.get("title_vi", "")
        difficulty = module.get("difficulty_level", "beginner")
        lesson_count = module.get("lesson_count", 0)

        print(f"\n   📦 Module {order}: {title}")

        # Module heading (H1)
        doc.add_heading(f"Module {order}: {title} — {title_vi}", level=1)

        # Module info
        diff_label = DIFFICULTY_LABELS.get(difficulty, difficulty)
        info_p = doc.add_paragraph()
        run = info_p.add_run(f"Difficulty: {diff_label}  •  Lessons: {lesson_count}")
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.italic = True

        # Separator
        sep_p = doc.add_paragraph()
        run = sep_p.add_run("─" * 60)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        # Process lessons in this module
        module_lessons = lessons_by_module.get(mid, [])
        for lesson in module_lessons:
            lid = lesson["id"]
            lorder = lesson["order_index"]
            ltitle = lesson["title"]
            ltitle_vi = lesson.get("title_vi", "")
            ldesc = lesson.get("description", "")
            content_json = lesson.get("content_json", "")

            total_lessons += 1
            print(f"      📖 Bài {lorder}: {ltitle}")

            # Lesson heading (H2)
            doc.add_heading(f"Bài {lorder}: {ltitle} — {ltitle_vi}", level=2)

            # Lesson description
            if ldesc:
                desc_p = doc.add_paragraph()
                add_inline_code_text(desc_p, ldesc)
                desc_p.paragraph_format.space_after = Pt(8)

            # Full lesson content from content_json
            render_content_sections(doc, content_json)

            # Quiz section
            add_quiz_section(doc, lid, quizzes, quiz_questions_by_quiz)

            # Keyword summary
            add_keyword_section(doc, lid, keywords)

            # Tech Lead Insights
            add_tech_lead_insights(doc, lid, ltitle)

            # Page break between lessons
            doc.add_page_break()

        # Page break between modules
        if module != modules[-1]:
            doc.add_page_break()

    # Save document
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"\n💾 Saving to {OUTPUT_FILE}...")
    doc.save(OUTPUT_FILE)

    file_size = os.path.getsize(OUTPUT_FILE)
    size_mb = file_size / (1024 * 1024)

    print(f"\n{'=' * 60}")
    print(f"✅ Document generated successfully!")
    print(f"   📄 File: {OUTPUT_FILE}")
    print(f"   📊 Size: {size_mb:.2f} MB")
    print(f"   📦 Modules: {len(modules)}")
    print(f"   📖 Lessons: {total_lessons}")
    print(f"   🔑 Keywords: {len(keywords)}")
    print(f"   📝 Quizzes: {len(quizzes)}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
