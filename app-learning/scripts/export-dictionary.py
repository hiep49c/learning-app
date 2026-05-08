#!/usr/bin/env python3
"""Export vocab-keywords.json to a dictionary-style .docx file."""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

SEED_PATH = Path(__file__).parent.parent / "assets" / "seed-data" / "vocab-keywords.json"
OUTPUT_PATH = Path(__file__).parent.parent / "English-Dictionary.docx"

LEVEL_ORDER = ["A1", "A2", "B1", "B2", "C1", "C2"]
LEVEL_NAMES = {
    "A1": "Beginner",
    "A2": "Elementary", 
    "B1": "Intermediate",
    "B2": "Upper-Intermediate",
    "C1": "Advanced",
    "C2": "Proficiency",
}


def deduplicate(keywords):
    """Merge duplicate words: keep longest definition/explanation, combine examples."""
    merged = {}
    for kw in keywords:
        name = kw["name"].lower().strip()
        if name not in merged:
            merged[name] = {
                "name": kw["name"].strip(),
                "definition": kw.get("definition", ""),
                "explanation": kw.get("explanation", ""),
                "code_example": kw.get("code_example", ""),
                "category": kw.get("category", "A1"),
                "usage_notes": kw.get("usage_notes", ""),
                "extra_examples": kw.get("extra_examples", []),
                "common_patterns": kw.get("common_patterns", ""),
                "_examples": set(),
            }
            if kw.get("code_example", "").strip():
                merged[name]["_examples"].add(kw["code_example"].strip())
        else:
            entry = merged[name]
            # Keep longest definition
            new_def = kw.get("definition", "")
            if len(new_def) > len(entry["definition"]):
                entry["definition"] = new_def
            # Keep longest explanation
            new_exp = kw.get("explanation", "")
            if len(new_exp) > len(entry["explanation"]):
                entry["explanation"] = new_exp
            # Keep lowest level (most basic)
            if LEVEL_ORDER.index(kw.get("category", "A1")) < LEVEL_ORDER.index(entry["category"]):
                entry["category"] = kw["category"]
            # Collect unique examples (max 3)
            ex = kw.get("code_example", "").strip()
            if ex and len(entry["_examples"]) < 3:
                entry["_examples"].add(ex)
            # Keep enriched fields if not already set
            if not entry["usage_notes"] and kw.get("usage_notes"):
                entry["usage_notes"] = kw["usage_notes"]
            if not entry["extra_examples"] and kw.get("extra_examples"):
                entry["extra_examples"] = kw["extra_examples"]
            if not entry["common_patterns"] and kw.get("common_patterns"):
                entry["common_patterns"] = kw["common_patterns"]

    # Finalize examples
    for entry in merged.values():
        examples = list(entry["_examples"])[:3]
        entry["code_example"] = "\n===\n".join(examples)
        del entry["_examples"]

    return list(merged.values())


def main():
    with open(SEED_PATH, "r", encoding="utf-8") as f:
        keywords = json.load(f)

    keywords = deduplicate(keywords)

    # Group by level, then sort alphabetically
    by_level = {}
    for kw in keywords:
        level = kw.get("category", "A1")
        by_level.setdefault(level, []).append(kw)

    for level in by_level:
        by_level[level].sort(key=lambda x: x["name"].lower())

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("English-Vietnamese Dictionary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Tổng hợp {len(keywords):,} từ vựng (không trùng lặp) — Từ A1 đến C2")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacing

    # Table of contents summary
    toc = doc.add_paragraph()
    toc.add_run("Mục lục:\n").bold = True
    for level in LEVEL_ORDER:
        count = len(by_level.get(level, []))
        if count:
            toc.add_run(f"  • {level} ({LEVEL_NAMES[level]}): {count:,} từ\n")

    doc.add_page_break()

    # Each level
    for level in LEVEL_ORDER:
        words = by_level.get(level, [])
        if not words:
            continue

        # Level heading
        doc.add_heading(f"{level} — {LEVEL_NAMES[level]} ({len(words):,} từ)", level=1)

        current_letter = ""

        for kw in words:
            first_letter = kw["name"][0].upper()
            if first_letter != current_letter:
                current_letter = first_letter
                letter_heading = doc.add_heading(current_letter, level=2)

            # Word entry
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)

            # Word name (bold, larger)
            word_run = p.add_run(kw["name"])
            word_run.bold = True
            word_run.font.size = Pt(12)

            # Definition
            definition = kw.get("definition", "")
            if definition:
                def_run = p.add_run(f"  —  {definition}")
                def_run.font.size = Pt(11)

            # Explanation (pronunciation, collocations, synonyms, antonyms)
            explanation = kw.get("explanation", "")
            if explanation:
                exp_p = doc.add_paragraph()
                exp_p.paragraph_format.left_indent = Cm(0.5)
                exp_p.paragraph_format.space_before = Pt(0)
                exp_p.paragraph_format.space_after = Pt(2)
                for line in explanation.split("\n"):
                    if line.strip():
                        run = exp_p.add_run(line.strip() + "\n")
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(80, 80, 80)

            # Example(s)
            example = kw.get("code_example", "")
            if example:
                ex_p = doc.add_paragraph()
                ex_p.paragraph_format.left_indent = Cm(0.5)
                ex_p.paragraph_format.space_before = Pt(0)
                ex_p.paragraph_format.space_after = Pt(2)
                for i, ex_block in enumerate(example.split("\n===\n")):
                    parts = ex_block.split("\n---\n")
                    en_part = parts[0].strip() if parts else ""
                    vi_part = parts[1].strip() if len(parts) > 1 else ""
                    if i > 0:
                        ex_p.add_run("\n")
                    if en_part:
                        r = ex_p.add_run(f"📝 {en_part}")
                        r.font.size = Pt(10)
                        r.italic = True
                    if vi_part:
                        ex_p.add_run("\n")
                        r2 = ex_p.add_run(f"    → {vi_part}")
                        r2.font.size = Pt(10)
                        r2.font.color.rgb = RGBColor(0, 100, 0)

            # Usage notes
            usage = kw.get("usage_notes", "")
            if usage:
                u_p = doc.add_paragraph()
                u_p.paragraph_format.left_indent = Cm(0.5)
                u_p.paragraph_format.space_before = Pt(0)
                u_p.paragraph_format.space_after = Pt(2)
                r = u_p.add_run(f"💡 {usage}")
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0, 80, 120)

            # Extra examples
            extra_ex = kw.get("extra_examples", [])
            if extra_ex:
                eex_p = doc.add_paragraph()
                eex_p.paragraph_format.left_indent = Cm(0.5)
                eex_p.paragraph_format.space_before = Pt(0)
                eex_p.paragraph_format.space_after = Pt(2)
                for i, ex_item in enumerate(extra_ex[:2]):
                    if not isinstance(ex_item, dict):
                        continue
                    en = ex_item.get("e") or ex_item.get("example_en", "")
                    vi = ex_item.get("v") or ex_item.get("example_vi", "")
                    if i > 0:
                        eex_p.add_run("\n")
                    if en:
                        r = eex_p.add_run(f"✏️ {en}")
                        r.font.size = Pt(9)
                        r.italic = True
                    if vi:
                        eex_p.add_run("\n")
                        r2 = eex_p.add_run(f"    → {vi}")
                        r2.font.size = Pt(9)
                        r2.font.color.rgb = RGBColor(0, 100, 0)

            # Common patterns
            patterns = kw.get("common_patterns", "")
            if patterns:
                if isinstance(patterns, list):
                    patterns = " | ".join(patterns)
                cp_p = doc.add_paragraph()
                cp_p.paragraph_format.left_indent = Cm(0.5)
                cp_p.paragraph_format.space_before = Pt(0)
                cp_p.paragraph_format.space_after = Pt(4)
                r = cp_p.add_run(f"🔗 Patterns: {patterns}")
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(100, 60, 0)

        doc.add_page_break()

    doc.save(str(OUTPUT_PATH))
    print(f"✅ Đã xuất từ điển: {OUTPUT_PATH}")
    print(f"   Tổng: {len(keywords):,} từ")


if __name__ == "__main__":
    main()
